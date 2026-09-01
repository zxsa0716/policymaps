#!/usr/bin/env python
"""원고(internal/서식2_본문.md) → 제출본 DOCX 조판.

한글(HWP)은 .docx 를 열 수 있으므로, 여기서 만든 파일을 한글에서 열고
'다른 이름으로 저장 → .hwp' 하면 된다.

**정본은 원고 .md 다.** 내용은 .md 를 고치고 이 스크립트를 다시 돌린다.
(예전에는 본문을 이 파일에 하드코딩해 두어 원고와 제출본이 갈라졌다.)

공식 작성요령(제출서류 양식.hwp 에서 OLE 파싱으로 추출)
    용지 A4 · 좌우여백 20mm · 위아래 10mm · 머리말/꼬리말 10mm
    줄간격 140% · 휴먼명조 15pt(본문) · 맑은 고딕 12pt(표·캡션)
    개조식으로 상세하게, 본문 15쪽 이내
    파란색 글씨(작성요령)는 제출 시 삭제 → 여기서는 애초에 넣지 않고,
    제목도 파란 계열을 피해 먹색으로 조판한다

원고 표기 규약
    ## 1  배경 및 개요       장
    ### □ 배경              □ 대분류
    **ㅇ 기획 목적**         ㅇ 중분류(1장)
    **◦ 데이터 활용 계획**    ◦ 중분류(2~5장)
    - 항목                  개조식 1수준
    · 항목                  개조식 2수준
    | a | b |              표(다음 줄 |---| 구분선)
    > _그림 N. 설명_         그림(FIGMAP 에서 경로를 찾는다)

사용:  python system/make_form2_docx.py
출력:  internal/서식2_본문.docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
FIG = ROOT / "docs" / "figures"
SHOT = ROOT / "docs" / "screenshots" / "doc"   # 문서용 크롭본(폭:높이 1.9:1)
if not SHOT.exists():
    SHOT = ROOT / "docs" / "screenshots"
SRC = ROOT / "internal" / "서식2_본문.md"
OUT = ROOT / "internal" / "서식2_본문.docx"

BODY_FONT = "휴먼명조"
TABLE_FONT = "맑은 고딕"
BODY_PT = 15          # 공식: 본문 휴먼명조 15pt
SMALL_PT = 12         # 공식: 표·캡션 맑은 고딕 12pt
LINE = 1.4            # 공식: 줄간격 140%
INK = RGBColor(0x1A, 0x1A, 0x1A)
TEXT_CM = 17.0        # A4 210mm - 좌우 여백 20mm

# 표 머리글 → 열 폭(cm). 합은 TEXT_CM(17.0). 자동 배분은 재현이 흔들려 쓰지 않는다.
WIDTHS = {
    ("기존 수단", "어디까지 제공하나", "어디서 멈추나", "본 과제가 더하는 것"):
        [3.4, 3.4, 4.4, 5.8],
    ("연번", "데이터명", "데이터 설명", "출처", "무상여부", "확보 실적"):
        [0.9, 2.5, 4.0, 4.2, 1.2, 4.2],
    ("데이터", "인증키·발급", "갱신 주기", "이용 조건"): [3.6, 4.6, 3.4, 5.4],
    ("가공", "처리 내용", "결과"): [2.0, 8.4, 6.6],
    ("단계", "처리 내용", "산출"): [1.8, 9.0, 6.2],
    ("노드", "수", "엣지", "수"): [5.0, 3.0, 5.2, 2.8],
    ("기법", "사양", "적용 결과"): [2.8, 8.0, 5.6],
    ("사분면", "지역 수", "구성"): [3.4, 2.0, 10.6],
    ("정책", "제정본", "지문 공유율", "시도 집중", "유형"): [3.2, 2.0, 3.0, 2.6, 5.2],
    ("모델", "분야 일치", "무작위 기준선", "배수"): [4.0, 3.4, 3.8, 3.0],
    ("검증 항목", "결과"): [3.8, 12.2],
    ("항목", "값", "항목", "값"): [4.0, 4.2, 3.8, 4.0],
    ("항목", "근거", "효과"): [3.4, 6.6, 6.0],
    ("한계", "수치", "후속"): [4.0, 7.0, 5.0],
    ("구분", "출처 및 주소"): [3.0, 13.0],
    ("구분", "출처"): [2.6, 13.4],
}

# 그림 번호 → (경로 목록, 폭cm). 두 장이면 나란히 배치한다.
FIGMAP = {
    "1": ([SHOT / "07_spatial.png"], 9.0),
    "2": ([SHOT / "02_map.png", SHOT / "15_ai_agent.png"], 6.7),
}


# --------------------------------------------------------------------------- #
# 조판 기초
# --------------------------------------------------------------------------- #
def set_run(run, *, size=BODY_PT, font=BODY_FONT, bold=False):
    run.font.size = Pt(size)
    run.font.name = font
    run.font.bold = bold
    run.font.color.rgb = INK
    # 한글 글꼴은 eastAsia 속성에 따로 지정해야 적용된다
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rFonts.set(qn(attr), font)


def para(doc, text="", *, size=BODY_PT, font=BODY_FONT, bold=False,
         indent=0.0, space_after=2, align=None, line=LINE):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    if indent:
        pf.left_indent = Cm(indent)
    if align is not None:
        p.alignment = align
    if text:
        set_run(p.add_run(text), size=size, font=font, bold=bold)
    return p


def chapter(doc, text):
    """장 제목 + 아래 실선. 공식 서식의 번호 박스를 대체한다."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    set_run(p.add_run(text), size=16, font=TABLE_FONT, bold=True)
    pPr = p._element.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "12")
    bot.set(qn("w:color"), "1A1A1A")
    bdr.append(bot)
    pPr.append(bdr)


def heading(doc, text, level):
    """level 1 = □ 대분류, level 2 = ㅇ/◦ 중분류."""
    para(doc, text if level == 1 else " " + text,
         size=BODY_PT, font=TABLE_FONT, bold=True,
         indent=0.0 if level == 1 else 0.15,
         space_after=1, line=1.15 if level == 1 else 1.2)


def bullet(doc, text, deep=False):
    para(doc, ("· " if deep else "- ") + text,
         size=BODY_PT, indent=0.8 if deep else 0.3, space_after=1)


def caption(doc, text):
    para(doc, text, size=SMALL_PT, font=TABLE_FONT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, line=1.1)


def figure(doc, paths, cap, width_cm):
    if len(paths) == 1:
        p = paths[0]
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(2)
        fp.paragraph_format.space_after = Pt(1)
        if p.exists():
            fp.add_run().add_picture(str(p), width=Cm(width_cm))
        else:
            set_run(fp.add_run(f"[그림 누락: {p.name}]"), size=SMALL_PT, font=TABLE_FONT)
    else:
        t = doc.add_table(rows=1, cols=len(paths))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True
        for cell, p in zip(t.rows[0].cells, paths):
            cell.text = ""
            pp = cell.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp.paragraph_format.space_after = Pt(0)
            if p.exists():
                pp.add_run().add_picture(str(p), width=Cm(width_cm))
            else:
                set_run(pp.add_run(f"[그림 누락: {p.name}]"), size=SMALL_PT, font=TABLE_FONT)
            tcPr = cell._tc.get_or_add_tcPr()
            bd = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                e = OxmlElement("w:" + side)
                e.set(qn("w:val"), "nil")
                bd.append(e)
            tcPr.append(bd)
    caption(doc, cap)


def _tight_cells(t, top=14, bottom=14):
    """표 셀 상하 여백 축소(Word 기본값의 약 1/4). 단위: twip."""
    mar = OxmlElement("w:tblCellMar")
    for side, v in (("top", top), ("bottom", bottom), ("left", 72), ("right", 72)):
        e = OxmlElement("w:" + side)
        e.set(qn("w:w"), str(v))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    t._tbl.tblPr.append(mar)


def table(doc, rows):
    """rows[0] 이 머리글. 열 폭은 내용 길이에 비례해 자동 배분한다."""
    ncol = len(rows[0])
    rows = [r + [""] * (ncol - len(r)) if len(r) < ncol else r[:ncol] for r in rows]
    widths = WIDTHS.get(tuple(rows[0]))
    if widths is None or len(widths) != ncol:
        span = [max(max(len(r[i]) for r in rows), 3) for i in range(ncol)]
        tot = sum(span)
        widths = [max(1.2, TEXT_CM * v / tot) for v in span]
        scale = TEXT_CM / sum(widths)
        widths = [w * scale for w in widths]
        print(f"    [폭 자동] {tuple(rows[0])}")

    t = doc.add_table(rows=1, cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(rows[0]):
        cell = t.rows[0].cells[i]
        cell.text = ""
        pp = cell.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.line_spacing = 1.05
        pp.paragraph_format.space_after = Pt(0)
        set_run(pp.add_run(h), size=SMALL_PT, font=TABLE_FONT, bold=True)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "EFEFEF")
        cell._tc.get_or_add_tcPr().append(shd)
    for r in rows[1:]:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = ""
            pp = cells[i].paragraphs[0]
            pp.paragraph_format.line_spacing = 1.05
            pp.paragraph_format.space_after = Pt(0)
            if i > 0 and v.replace(",", "").replace(".", "").replace("%", "").isdigit():
                pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_run(pp.add_run(v), size=SMALL_PT, font=TABLE_FONT)
    for i, w in enumerate(widths):
        for row in t.rows:
            row.cells[i].width = Cm(w)
    _tight_cells(t)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(0)
    sp.paragraph_format.line_spacing = 0.5
    return t


def add_page_number(section):
    """꼬리말에 '- N -' 쪽번호 필드."""
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("- "), size=SMALL_PT, font=TABLE_FONT)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(SMALL_PT * 2)))
    rPr.append(sz)
    r.append(rPr)
    fld.append(r)
    p._p.append(fld)
    set_run(p.add_run(" -"), size=SMALL_PT, font=TABLE_FONT)


# --------------------------------------------------------------------------- #
# 원고 파싱
# --------------------------------------------------------------------------- #
RE_FIG = re.compile(r"^>\s*_그림\s*([0-9]+)(?:·[0-9]+)?\.\s*(.*?)_\s*$")
RE_TCAP = re.compile(r"^표\s*[0-9]+\..*$")


def parse(md: str):
    """원고 → 조판 명령 목록."""
    out = []
    lines = md.split("\n")
    i = lines.index("---") + 1 if "---" in lines else 0
    while i < len(lines):
        s = lines[i].strip()
        if not s:
            i += 1
            continue
        m = RE_FIG.match(s)
        if s.startswith("## "):
            out.append(("chapter", s[3:].strip()))
        elif s.startswith("### "):
            out.append(("head1", s[4:].strip()))
        elif s.startswith("**") and s.endswith("**"):
            out.append(("head2", s[2:-2].strip()))
        elif m:
            out.append(("figure", m.group(1), m.group(2)))
        elif (s.startswith("|") and i + 1 < len(lines)
              and set(lines[i + 1].strip()) <= set("|-: ")
              and "-" in lines[i + 1]):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if set("".join(cells)) <= set("-: "):
                    i += 1
                    continue
                rows.append(cells)
                i += 1
            out.append(("table", rows))
            continue
        elif RE_TCAP.match(s):
            out.append(("tcap", s))
        elif s.startswith("- "):
            out.append(("bullet", s[2:].strip(), False))
        elif s.startswith("· "):
            out.append(("bullet", s[2:].strip(), True))
        else:
            out.append(("text", s))
        i += 1
    return out


def build() -> int:
    if not SRC.exists():
        print(f"원고가 없다: {SRC}")
        return 1
    ops = parse(SRC.read_text(encoding="utf-8"))

    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Mm(210), Mm(297)          # A4
    s.left_margin = s.right_margin = Mm(20)
    s.top_margin = s.bottom_margin = Mm(10)
    s.header_distance = s.footer_distance = Mm(10)
    add_page_number(s)

    nfig = 0
    ntitle = 0
    for op in ops:
        kind = op[0]
        if kind == "text":
            ntitle += 1
            if ntitle == 1:      # 표지 제목
                para(doc, op[1], size=21, font=TABLE_FONT, bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, line=1.15)
            elif ntitle == 2:    # 부제
                para(doc, op[1], size=SMALL_PT + 1, font=TABLE_FONT,
                     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, line=1.2)
            else:
                para(doc, op[1])
        elif kind == "chapter":
            chapter(doc, op[1])
        elif kind == "head1":
            heading(doc, op[1], 1)
        elif kind == "head2":
            heading(doc, op[1], 2)
        elif kind == "bullet":
            bullet(doc, op[1], op[2])
        elif kind == "table":
            table(doc, op[1])
        elif kind == "tcap":
            caption(doc, op[1])
        elif kind == "figure":
            paths, w = FIGMAP.get(op[1], ([], 9.0))
            if not paths:
                continue
            nfig += len(paths)
            no = op[1] if len(paths) == 1 else f"{op[1]}·{int(op[1]) + len(paths) - 1}"
            figure(doc, paths, f"그림 {no}. {op[2]}", w)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"  → {OUT}  ({OUT.stat().st_size // 1024}KB, 그림 {nfig}장)")
    return 0


if __name__ == "__main__":
    raise SystemExit(build())
